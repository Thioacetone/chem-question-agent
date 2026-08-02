"""
Word导出服务 v2.0 - 严格仿照江苏高考真题排版格式
基于2021-2024年真题及模拟卷格式分析
支持合成路线图嵌入和结构式渲染
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re

from structure_renderer import renderer


class ExportService:
    """文档导出服务 - 仿高考真题格式"""

    # ================================================================
    # 真题排版参数（基于2021-2024江苏高考真题+模拟卷分析）
    # ================================================================
    BODY_FONT = '宋体'          # 正文字体
    BODY_SIZE = Pt(10.5)        # 五号 = 10.5pt
    WESTERN_FONT = 'Times New Roman'  # 西文字体
    LINE_SPACING = 1.5          # 1.5倍行距
    PAGE_WIDTH = Cm(21)         # A4
    PAGE_HEIGHT = Cm(29.7)
    MARGIN_TOP = Cm(2.54)
    MARGIN_BOTTOM = Cm(2.54)
    MARGIN_LEFT = Cm(3.17)
    MARGIN_RIGHT = Cm(3.17)

    @staticmethod
    def _set_run_font(run, font_name=None, size=None, bold=False, color=None):
        """设置run的字体属性"""
        run.font.name = font_name or ExportService.BODY_FONT
        run.font.size = size or ExportService.BODY_SIZE
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        # 设置西文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name or ExportService.BODY_FONT)
        rFonts.set(qn('w:ascii'), ExportService.WESTERN_FONT)
        rFonts.set(qn('w:hAnsi'), ExportService.WESTERN_FONT)

    @staticmethod
    def _add_paragraph(doc, text, indent=False, font_size=None, bold=False, alignment=None):
        """添加一个标准格式的段落"""
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing = ExportService.LINE_SPACING
        if indent:
            pf.first_line_indent = Cm(0.74)  # 两个字符缩进
        if alignment is not None:
            pf.alignment = alignment

        if text:
            run = para.add_run(text)
            ExportService._set_run_font(run, size=font_size, bold=bold)
        return para

    @staticmethod
    def _embed_route_diagram(doc, route_data: dict):
        """
        嵌入合成路线流程图（仿照高考真题中的路线图）
        
        Args:
            doc: Word Document对象
            route_data: 合成路线数据 {"steps": [{"reactant": "苯", "reagent": "HNO3/H2SO4", "product": "硝基苯", "reaction_type": "硝化"}, ...]}
        """
        steps = route_data.get("steps", [])
        if not steps:
            return

        # 构建路线图步骤数据
        diagram_steps = []
        for i, step in enumerate(steps):
            # 尝试将名称转为SMILES
            reactant_name = step.get("reactant", "")
            reactant_smiles = renderer.name_to_smiles(reactant_name)
            if not reactant_smiles:
                # 尝试直接使用SMILES
                raw = step.get("reactant", "")
                if renderer.smiles_to_mol(raw):
                    reactant_smiles = raw

            # 第一个化合物只需要reactant
            if i == 0:
                diagram_steps.append({
                    "smiles": reactant_smiles or "",
                    "label": chr(65 + i),  # A, B, C...
                    "name": reactant_name,
                })

            # 产物
            product_name = step.get("product", "")
            product_smiles = renderer.name_to_smiles(product_name)
            if not product_smiles:
                raw = step.get("product", "")
                if renderer.smiles_to_mol(raw):
                    product_smiles = raw

            reagent = step.get("reagent", "")
            diagram_steps.append({
                "smiles": product_smiles or "",
                "label": chr(65 + i + 1),
                "reagent": reagent,
                "name": product_name,
            })

        if len(diagram_steps) < 2:
            return

        try:
            # 生成路线图PNG
            png_bytes = renderer.render_route_diagram_png(diagram_steps)
            if not png_bytes:
                return

            # 嵌入到Word文档
            image_stream = io.BytesIO(png_bytes)
            
            # 图片段落居中
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
            
            run = img_para.add_run()
            # 计算合适的图片尺寸
            max_width = Cm(14)  # 最大宽度14cm
            run.add_picture(image_stream, width=max_width)
            
            # 空行
            ExportService._add_paragraph(doc, "")
        except Exception as e:
            # 渲染失败不阻塞导出
            print(f"路线图渲染失败: {e}")

    @staticmethod
    def _extract_route_text(route_data: dict) -> str:
        """从路线数据生成文字版路线描述（试题中'其合成路线如下：'后面的文字）"""
        steps = route_data.get("steps", [])
        if not steps:
            return ""
        
        # 生成 A→B→C→... 格式
        compounds = []
        for i, step in enumerate(steps):
            if i == 0:
                compounds.append(f"A（{step.get('reactant', '')}）")
            compounds.append(f"{chr(65 + i + 1)}（{step.get('product', '')}）")
        
        return " → ".join(compounds)

    @staticmethod
    def _add_question_group(doc, questions, answers=None):
        """添加小题组，仿照真题格式"""
        for q in questions:
            q_num = q.get("number", "?")
            q_content = q.get("content", "")
            q_type = q.get("type", "")

            # 去除content中可能已有的编号前缀，如 "(1)" "1." "1、" 等
            clean_content = re.sub(r'^[（(]?\d+[）).．、]\s*', '', q_content.strip())

            # 小题正文
            q_para = doc.add_paragraph()
            q_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
            q_para.paragraph_format.first_line_indent = Cm(0.74)

            # 编号用全角括号
            num_run = q_para.add_run(f"（{q_num}）")
            ExportService._set_run_font(num_run)

            content_run = q_para.add_run(clean_content)
            ExportService._set_run_font(content_run)

            # 同分异构体：条件用①②单独段落
            if "同分异构" in q_type:
                # 尝试从content中提取条件（如果content包含①②③）
                conditions = re.findall(r'[①②③][^①②③]+', clean_content)
                if conditions:
                    # 移除条件部分，只保留问题主干
                    main_text = clean_content
                    for cond in conditions:
                        main_text = main_text.replace(cond, '')
                    main_text = re.sub(r'\s+', ' ', main_text).strip()
                    # 更新content_run
                    content_run.text = main_text

                    # 添加条件段落
                    for cond in conditions:
                        cond_para = doc.add_paragraph()
                        cond_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
                        cond_para.paragraph_format.first_line_indent = Cm(0.74)
                        cond_run = cond_para.add_run(cond.strip())
                        ExportService._set_run_font(cond_run)

            # 合成路线设计：已知信息单独段落
            if "合成路线" in q_type and "已知" in clean_content:
                # 拆分已知信息
                known_match = re.match(r'(已知[：:][^。]*[。]?)(.*)', clean_content)
                if known_match:
                    known_text = known_match.group(1).strip()
                    rest_text = known_match.group(2).strip()

                    # 已知信息单独段落
                    content_run.text = rest_text
                    known_para = doc.add_paragraph()
                    known_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
                    known_para.paragraph_format.first_line_indent = Cm(0.74)
                    known_run = known_para.add_run(known_text)
                    ExportService._set_run_font(known_run)

        # 答案部分
        if answers:
            for ans in answers:
                ans_num = ans.get("number", "?")
                ans_content = ans.get("content", "")

                ans_para = doc.add_paragraph()
                ans_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
                ans_para.paragraph_format.first_line_indent = Cm(0.74)

                num_run = ans_para.add_run(f"（{ans_num}）")
                ExportService._set_run_font(num_run)

                content_run = ans_para.add_run(ans_content)
                ExportService._set_run_font(content_run)

                # 踩分点
                scoring = ans.get("scoring_points", [])
                if scoring:
                    for sp in scoring:
                        sp_para = doc.add_paragraph()
                        sp_para.paragraph_format.line_spacing = ExportService.LINE_SPACING
                        sp_para.paragraph_format.left_indent = Cm(1.5)
                        sp_run = sp_para.add_run(f"【踩分点】{sp}")
                        ExportService._set_run_font(sp_run, size=Pt(9), color=RGBColor(180, 0, 0))

    @staticmethod
    def export_to_docx(question_data: dict, include_answer: bool = True) -> bytes:
        """
        导出为Word文档，仿照江苏高考真题格式

        Args:
            question_data: 命题数据
            include_answer: 是否包含答案（教师版=True，学生版=False）
        """
        doc = Document()

        # --- 页面设置 ---
        for section in doc.sections:
            section.page_width = ExportService.PAGE_WIDTH
            section.page_height = ExportService.PAGE_HEIGHT
            section.top_margin = ExportService.MARGIN_TOP
            section.bottom_margin = ExportService.MARGIN_BOTTOM
            section.left_margin = ExportService.MARGIN_LEFT
            section.right_margin = ExportService.MARGIN_RIGHT

        # --- 默认样式 ---
        style = doc.styles['Normal']
        style.font.name = ExportService.BODY_FONT
        style.font.size = ExportService.BODY_SIZE
        style.paragraph_format.line_spacing = ExportService.LINE_SPACING

        # ================================================================
        # 学生版（空白卷）格式
        # ================================================================

        # 题干：仿照真题 "15. 化合物F是[用途身份]，其合成路线如下："
        stem = question_data.get("stem", "")
        if stem:
            # 确保题干以题号开头
            stem_clean = stem.strip()
            if not re.match(r'^\d+[\.．]', stem_clean):
                stem_clean = f"15. {stem_clean}"

            stem_para = ExportService._add_paragraph(doc, stem_clean)
            stem_para.paragraph_format.first_line_indent = Cm(0.74)

        # 嵌入合成路线图（仿照真题中的路线图）
        raw_route = question_data.get("raw_route", {})
        if raw_route and raw_route.get("steps"):
            ExportService._embed_route_diagram(doc, raw_route)
        else:
            ExportService._add_paragraph(doc, "")

        # 小题
        questions = question_data.get("questions", [])
        ExportService._add_question_group(doc, questions)

        # ================================================================
        # 教师版：追加参考答案和解析
        # ================================================================
        if include_answer:
            doc.add_page_break()

            # 标题
            title_para = ExportService._add_paragraph(
                doc, "参考答案与解析", bold=True,
                font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
            ExportService._add_paragraph(doc, "")

            answers = question_data.get("answers", [])
            ExportService._add_question_group(doc, questions, answers)

            ExportService._add_paragraph(doc, "")

            # 解析
            if "analysis" in question_data and question_data["analysis"]:
                analysis_heading = ExportService._add_paragraph(
                    doc, "【详细解析】", bold=True, font_size=Pt(12)
                )
                analysis_text = question_data["analysis"]
                analysis_para = ExportService._add_paragraph(doc, analysis_text, indent=True)

            ExportService._add_paragraph(doc, "")

            # 评分标准
            scoring_heading = ExportService._add_paragraph(
                doc, "【评分标准】", bold=True, font_size=Pt(12)
            )
            scoring_lines = [
                "1. 官能团名称：每个正确官能团得1分，错别字扣1分",
                "2. 结构简式：完全正确得满分，结构不规范扣1分",
                "3. 反应类型：答对得满分，错别字扣1分",
                "4. 同分异构体：写出正确结构得满分（3分），多写不扣分，写错不倒扣",
                "5. 合成路线设计：每一步正确得1分，整体逻辑正确得1-2分，共5分",
            ]
            for line in scoring_lines:
                ExportService._add_paragraph(doc, line, indent=True)

        # 保存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


export_service = ExportService()